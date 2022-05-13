#!/usr/bin/env python
# -*- coding:utf-8 -*-

#pip3 install python-nmap && pip3 install python-docx
#python3 ./nmap_scan.py IP/掩码 端口 或 python3 ./nmap_scan.py IP/掩码 端口 "参数"
#在当前目录下生成名为nmap.docx的文件
import nmap
import sys
import docx

def test():
        ip = sys.argv[1]
        port_input = sys.argv[2]
        nm = nmap.PortScanner()
        if len(sys.argv) == 3:
                nm.scan(ip, port_input)
        else:
                cmd = sys.argv[3]
                nm.scan(ip, port_input, cmd)
#       print(nm.csv())
        file=docx.Document()
        for host in nm.all_hosts():
                file.add_paragraph('——————————————————————————————————————')
                var_1 = 'Host:' + host + ' ' + nm[host].state()
                file.add_paragraph(var_1)
#               print('Host : %s (%s)' % (host, nm[host].hostname()))
                var_2 = 'State :' + nm[host].state()
                file.add_paragraph(var_2)
#               print('State : %s' % nm[host].state())
                for proto in nm[host].all_protocols():
#                       print('----------')
                        var_3 = 'Protocol :' + proto
                        file.add_paragraph(var_3)
#                       print('Protocol : %s' % proto)
                        lport = nm[host][proto].keys()
#                       lport.sort()
                        for port in lport:
                                var_4 = 'port: ' + str(port) + ' ' + 'state: ' + nm[host][proto][port]['state'] + ' ' +'name: ' + nm[host][proto][port]['name'] + ' ' + 'product:' + ' '+  nm[host][proto][port]['product'] + ' ' + nm[host][proto][port]['version']
                                file.add_paragraph(var_4)
#                               print(var_4)
#                               print ('port : %s\tstate : %s\t%s' % (port, nm[host][proto][port]['state'] , nm[host][proto][port]['name']))
        file.save("./nmap.docx")
        print("扫描完成")
if __name__ == '__main__':
        test()
